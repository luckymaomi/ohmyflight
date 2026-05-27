#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
人员结构报告填充脚本。

输入人员信息 Excel 和人员结构报告 docx，按人员结构统计口径填写飞行部前 9 张统计表，
并另存为新 docx，不覆盖原文件。
"""

from __future__ import annotations

import argparse
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Iterable

try:
    from openpyxl import load_workbook
except ImportError:  # pragma: no cover
    load_workbook = None

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


REQUIRED_HEADERS = [
    "姓名",
    "技术信息",
    "RAMA",
    "REUO",
    "RWAS",
    "EAMA",
    "EEUO",
    "EWAS",
    "原单位",
    "检查员资格",
    "行政职务",
]

QUALIFICATION_CODES = [
    "RAMA",
    "REUO",
    "RWAS",
    "RSEA",
    "EAMA",
    "EEUO",
    "EWAS",
    "ESEA",
    "RANC",
    "RORD",
    "RJFK",
    "RLAX",
    "RNLU",
]

ORIGIN_LABELS = [
    "飞行/总队 777",
    "飞行/总队 737",
    "飞行/总队 320",
    "飞行/总队 909",
    "湖南",
    "湖北",
    "新疆",
    "河南",
    "西安",
    "重庆",
    "汕头",
    "珠海",
    "广西",
    "海南",
    "上海",
]

FIXED_GROUND_COUNT = 29
MONTH_PATTERN = re.compile(r"^(1[0-2]|[1-9])月$")


@dataclass
class PersonnelRecord:
    employee_id: str
    name: str
    tech_info: str
    origin: str
    inspector_qualification: str
    management_role: str
    qualifications: dict[str, bool]


@dataclass
class StatItem:
    label: str
    count: int
    denominator: int
    percent: str
    rule: str


@dataclass
class StatSection:
    title: str
    denominator_label: str
    items: list[StatItem]


@dataclass
class StatResult:
    total_people: int
    registered_crew_count: int
    ground_count: int
    sections: list[StatSection]
    warnings: list[str]


def normalize_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def normalize_header(value) -> str:
    return re.sub(r"\s+", "", normalize_text(value))


def has_value(value) -> bool:
    return normalize_text(value) != ""


def tech_label(record: PersonnelRecord) -> str:
    parts = re.split(r"[:：]", record.tech_info, maxsplit=1)
    return normalize_text(parts[1] if len(parts) > 1 else record.tech_info)


def is_teacher(record: PersonnelRecord) -> bool:
    return "飞行教员" in tech_label(record)


def is_transfer_captain(record: PersonnelRecord) -> bool:
    return tech_label(record) == "划转机长"


def is_transfer_first_officer(record: PersonnelRecord) -> bool:
    return tech_label(record) == "划转副驾驶"


def is_captain(record: PersonnelRecord) -> bool:
    label = tech_label(record)
    return "机长" in label and "飞行教员" not in label and "划转" not in label


def is_regular_first_officer(record: PersonnelRecord) -> bool:
    label = tech_label(record)
    return "副驾驶" in label and "划转" not in label


def is_registered_crew(record: PersonnelRecord) -> bool:
    return is_teacher(record) or is_captain(record) or is_regular_first_officer(record)


def is_captain_or_above(record: PersonnelRecord) -> bool:
    return is_teacher(record) or is_captain(record) or is_transfer_captain(record)


def is_first_officer_group(record: PersonnelRecord) -> bool:
    return is_regular_first_officer(record) or is_transfer_first_officer(record)


def is_structure_crew(record: PersonnelRecord) -> bool:
    return is_captain_or_above(record) or is_first_officer_group(record)


def has_qualification(record: PersonnelRecord, code: str) -> bool:
    return bool(record.qualifications.get(code))


def is_local(record: PersonnelRecord) -> bool:
    return record.origin.startswith("总队") or record.origin == "777返聘"


def is_inspector(record: PersonnelRecord) -> bool:
    return record.inspector_qualification in {"公司检查员", "委任代表"}


def is_line_captain(record: PersonnelRecord) -> bool:
    label = tech_label(record)
    is_captain_level = (
        "飞行教员" in label
        or "A类机长" in label
        or "B类机长" in label
        or "C类机长" in label
    )
    return (
        is_captain_level
        and "Z类机长" not in label
        and not has_qualification(record, "RAMA")
        and not has_qualification(record, "REUO")
        and not has_qualification(record, "RWAS")
    )


def percent(count: int, denominator: int) -> str:
    if denominator == 0:
        return "0%"
    return f"{round(count / denominator * 100)}%"


def make_item(label: str, count: int, denominator: int, rule: str = "") -> StatItem:
    return StatItem(label=label, count=count, denominator=denominator, percent=percent(count, denominator), rule=rule)


def count(records: Iterable[PersonnelRecord], predicate: Callable[[PersonnelRecord], bool]) -> int:
    return sum(1 for record in records if predicate(record))


def format_people(records: list[PersonnelRecord]) -> str:
    values = []
    for record in records:
        identity = " ".join(part for part in [record.employee_id, record.name] if part)
        values.append(identity or tech_label(record))
    return "、".join(values)


def make_closed_items(
    records: list[PersonnelRecord],
    denominator: int,
    definitions: list[tuple[str, Callable[[PersonnelRecord], bool], str]],
    other_rule_prefix: str,
) -> list[StatItem]:
    matched: set[int] = set()
    items: list[StatItem] = []
    for label, predicate, rule in definitions:
        matched_records = [record for record in records if predicate(record)]
        for record in matched_records:
            matched.add(id(record))
        items.append(make_item(label, len(matched_records), denominator, rule))
    other_records = [record for record in records if id(record) not in matched]
    if other_records:
        items.append(make_item("其他", len(other_records), denominator, f"{other_rule_prefix}：{format_people(other_records)}。"))
    return items


def make_combo_items(
    records: list[PersonnelRecord],
    denominator: int,
    prefix: str,
    labels: dict[str, str],
    rule_prefix: str,
) -> list[StatItem]:
    north = f"{prefix}AMA"
    europe = f"{prefix}EUO"
    west = f"{prefix}WAS"

    def combo_count(expected_north: bool, expected_europe: bool, expected_west: bool) -> int:
        return count(
            records,
            lambda record: has_qualification(record, north) == expected_north
            and has_qualification(record, europe) == expected_europe
            and has_qualification(record, west) == expected_west,
        )

    return [
        make_item("美+欧+西亚", combo_count(True, True, True), denominator, f"{rule_prefix}：同时具备北美、欧洲、西亚。"),
        make_item("美+欧", combo_count(True, True, False), denominator, f"{rule_prefix}：具备北美、欧洲，不具备西亚。"),
        make_item("美+西亚", combo_count(True, False, True), denominator, f"{rule_prefix}：具备北美、西亚，不具备欧洲。"),
        make_item("欧+西亚", combo_count(False, True, True), denominator, f"{rule_prefix}：具备欧洲、西亚，不具备北美。"),
        make_item(labels["north_only"], combo_count(True, False, False), denominator, f"{rule_prefix}：只具备北美。"),
        make_item(labels["europe_only"], combo_count(False, True, False), denominator, f"{rule_prefix}：只具备欧洲。"),
        make_item(labels["west_only"], combo_count(False, False, True), denominator, f"{rule_prefix}：只具备西亚。"),
        make_item(labels["none"], combo_count(False, False, False), denominator, f"{rule_prefix}：北美、欧洲、西亚均不具备。"),
    ]


def build_captain_route_items(records: list[PersonnelRecord], denominator: int) -> list[StatItem]:
    combo_items = [
        item
        for item in make_combo_items(
            records,
            denominator,
            "R",
            {
                "north_only": "仅北美带队",
                "europe_only": "仅欧洲带队",
                "west_only": "仅西亚带队",
                "none": "无美欧西亚单飞",
            },
            "RAMA/REUO/RWAS 单飞资格",
        )
        if item.label != "无美欧西亚单飞"
    ]
    matched: set[int] = set()
    for record in records:
        has_any_single_flight = (
            has_qualification(record, "RAMA")
            or has_qualification(record, "REUO")
            or has_qualification(record, "RWAS")
        )
        if has_any_single_flight or is_line_captain(record) or tech_label(record) == "Z类机长":
            matched.add(id(record))
    unmatched = [record for record in records if id(record) not in matched]
    return [
        *combo_items,
        make_item("航线机长", count(records, is_line_captain), denominator, "B类及以上、无RAMA/REUO/RWAS单飞资格、且不是Z类机长。"),
        make_item("左座带飞", count(records, lambda record: tech_label(record) == "Z类机长"), denominator, "Z类机长。"),
        make_item(
            "其他",
            len(unmatched),
            denominator,
            f"不属于上述航线资格分类，需人工核对：{format_people(unmatched)}。" if unmatched else "上述航线资格分类已覆盖全部人员。",
        ),
    ]


def build_captain_level_items(records: list[PersonnelRecord], denominator: int) -> list[StatItem]:
    return [
        make_item("检查员", count(records, is_inspector), denominator, "检查员资格为公司检查员或委任代表。"),
        *make_closed_items(
            records,
            denominator,
            [
                ("C类教员", lambda record: tech_label(record) == "飞行教员C", "技术信息为飞行教员C。"),
                ("B类教员", lambda record: tech_label(record) == "飞行教员B", "技术信息为飞行教员B。"),
                ("A类教员", lambda record: tech_label(record) == "飞行教员A", "技术信息为飞行教员A。"),
                ("D类机长", lambda record: tech_label(record) == "D类机长", "技术信息为D类机长。"),
                ("C类机长", lambda record: tech_label(record) == "C类机长", "技术信息为C类机长。"),
                ("B类机长", lambda record: tech_label(record) == "B类机长", "技术信息为B类机长。"),
                ("Z类机长", lambda record: tech_label(record) == "Z类机长", "技术信息为Z类机长。"),
                ("在训机长", is_transfer_captain, "划转机长。"),
            ],
            "未落入教员/机长等级分类，需人工核对",
        ),
    ]


def build_first_officer_level_items(records: list[PersonnelRecord], denominator: int) -> list[StatItem]:
    return make_closed_items(
        records,
        denominator,
        [
            ("D类副驾驶", lambda record: tech_label(record) == "D类副驾驶", "技术信息为D类副驾驶。"),
            ("C类副驾驶", lambda record: tech_label(record) == "C类副驾驶", "技术信息为C类副驾驶。"),
            ("B类副驾驶", lambda record: tech_label(record) == "B类副驾驶", "技术信息为B类副驾驶。"),
            ("A类副驾驶", lambda record: tech_label(record) in {"A1类副驾驶", "A2类副驾驶"}, "技术信息为A1类副驾驶或A2类副驾驶。"),
            ("转机型副驾驶", is_transfer_first_officer, "划转副驾驶。"),
        ],
        "未落入副驾驶等级分类，需人工核对",
    )


def map_origin(origin: str) -> str:
    normalized = normalize_text(origin)
    mapping = {
        "总队777": "飞行/总队 777",
        "777返聘": "飞行/总队 777",
        "总队737": "飞行/总队 737",
        "总队320": "飞行/总队 320",
        "总队909": "飞行/总队 909",
        "湖南分公司": "湖南",
        "湖北分公司": "湖北",
        "新疆分公司": "新疆",
        "新疆分公司（借）": "新疆",
        "河南分公司": "河南",
        "西安分公司": "西安",
        "重庆航空": "重庆",
        "汕头分公司": "汕头",
        "珠海分公司": "珠海",
        "广西分公司": "广西",
        "海南分公司": "海南",
        "上海分公司（借）": "上海",
    }
    return mapping.get(normalized, normalized or "未识别")


def find_header_row(rows: list[list]) -> int:
    for index, row in enumerate(rows):
        headers = [normalize_header(value) for value in row]
        if sum(1 for header in REQUIRED_HEADERS if header in headers) >= 5:
            return index
    raise ValueError(f"未识别到人员信息表表头，至少需要包含：{'、'.join(REQUIRED_HEADERS)}")


def parse_excel(path: Path, sheet_name: str | None = None) -> list[PersonnelRecord]:
    if load_workbook is None:
        raise RuntimeError("缺少 openpyxl，请先运行：pip install openpyxl python-docx")

    workbook = load_workbook(path, data_only=True, read_only=True)
    if sheet_name:
        sheets = [workbook[sheet_name]]
    else:
        sheets = [workbook[name] for name in workbook.sheetnames]

    last_error: Exception | None = None
    for sheet in sheets:
        try:
            return parse_excel_sheet(sheet)
        except ValueError as exc:
            last_error = exc
            continue

    if last_error:
        raise last_error
    return []


def parse_excel_sheet(sheet) -> list[PersonnelRecord]:
    rows = [[cell for cell in row] for row in sheet.iter_rows(values_only=True)]
    header_row_index = find_header_row(rows)
    header_row = rows[header_row_index]
    header_map = {normalize_header(value): index for index, value in enumerate(header_row) if normalize_header(value)}
    missing_headers = [header for header in REQUIRED_HEADERS if header not in header_map]
    if missing_headers:
        raise ValueError(f"人员信息表缺少必要表头：{'、'.join(missing_headers)}")

    def value_by_header(row: list, header: str):
        index = header_map.get(header)
        if index is None or index >= len(row):
            return None
        return row[index]

    records: list[PersonnelRecord] = []
    for row in rows[header_row_index + 1 :]:
        employee_id = normalize_text(value_by_header(row, "员工号"))
        name = normalize_text(value_by_header(row, "姓名"))
        tech_info = normalize_text(value_by_header(row, "技术信息"))
        if not employee_id and not name and not tech_info:
            continue
        qualifications = {code: has_value(value_by_header(row, code)) for code in QUALIFICATION_CODES}
        records.append(
            PersonnelRecord(
                employee_id=employee_id,
                name=name,
                tech_info=tech_info,
                origin=normalize_text(value_by_header(row, "原单位")),
                inspector_qualification=normalize_text(value_by_header(row, "检查员资格")),
                management_role=normalize_text(value_by_header(row, "行政职务")),
                qualifications=qualifications,
            )
        )
    return records


def calculate(records: list[PersonnelRecord]) -> StatResult:
    registered_crew = [record for record in records if is_registered_crew(record)]
    structure_crew = [record for record in records if is_structure_crew(record)]
    captain_base = [record for record in records if is_teacher(record) or is_captain(record)]
    captain_with_training = [record for record in records if is_captain_or_above(record)]
    first_officer_base = [record for record in records if is_regular_first_officer(record)]
    first_officer_with_transfer = [record for record in records if is_first_officer_group(record)]

    registered_crew_count = len(registered_crew)
    structure_crew_count = len(structure_crew)
    ground_count = FIXED_GROUND_COUNT
    total_people = registered_crew_count + ground_count

    sections = [
        StatSection(
            "总人数及空地人员占比",
            f"{total_people}人",
            [
                make_item("总人数", total_people, total_people, "已注册空勤人员加固定地面人员29人。"),
                make_item("空勤人员（已注册人员）", registered_crew_count, total_people, "飞行教员、非划转机长、非划转副驾驶。"),
                make_item("地面人员", ground_count, total_people, "固定按29人统计。"),
            ],
        ),
        StatSection(
            "飞行管理人员占比",
            f"{structure_crew_count}人",
            [
                make_item("管理人员", count(structure_crew, lambda record: bool(record.management_role)), structure_crew_count, "行政职务非空。"),
                make_item("非管理人员", count(structure_crew, lambda record: not record.management_role), structure_crew_count, "行政职务为空。"),
            ],
        ),
        StatSection(
            "教员、机长、副驾驶占比",
            f"{structure_crew_count}人",
            [
                make_item("教员", count(structure_crew, is_teacher), structure_crew_count, "技术信息包含飞行教员。"),
                make_item("机长", count(structure_crew, lambda record: is_captain(record) or is_transfer_captain(record)), structure_crew_count, "非教员机长，含划转机长。"),
                make_item("副驾驶", count(structure_crew, lambda record: is_regular_first_officer(record) or is_transfer_first_officer(record)), structure_crew_count, "副驾驶，含划转副驾驶。"),
            ],
        ),
        StatSection("机长含以上各级别占比", f"{len(captain_with_training)}人", build_captain_level_items(captain_with_training, len(captain_with_training))),
        StatSection("机长航线资格占比", f"{len(captain_base)}人", build_captain_route_items(captain_base, len(captain_base))),
        StatSection(
            "机长报务占比",
            f"{len(captain_base)}人",
            make_combo_items(
                captain_base,
                len(captain_base),
                "E",
                {"north_only": "单美洲报务", "europe_only": "单欧洲报务", "west_only": "单西亚报务", "none": "无报务"},
                "EAMA/EEUO/EWAS 英语通信资格",
            ),
        ),
        StatSection("副驾驶级别占比", f"{len(first_officer_with_transfer)}人", build_first_officer_level_items(first_officer_with_transfer, len(first_officer_with_transfer))),
        StatSection(
            "副驾驶报务占比",
            f"{len(first_officer_base)}人",
            make_combo_items(
                first_officer_base,
                len(first_officer_base),
                "E",
                {"north_only": "单美洲报务", "europe_only": "单欧洲报务", "west_only": "单西亚报务", "none": "无报务"},
                "EAMA/EEUO/EWAS 英语通信资格",
            ),
        ),
        StatSection(
            "人员居住情况",
            f"{len(captain_with_training)}人 / {len(first_officer_with_transfer)}人",
            [
                make_item("机长本地居住", count(captain_with_training, is_local), len(captain_with_training), "原单位以总队开头或等于777返聘。"),
                make_item("机长异地居住", count(captain_with_training, lambda record: not is_local(record)), len(captain_with_training), "除本地外均为异地。"),
                make_item("副驾驶本地居住", count(first_officer_with_transfer, is_local), len(first_officer_with_transfer), "原单位以总队开头或等于777返聘。"),
                make_item("副驾驶异地居住", count(first_officer_with_transfer, lambda record: not is_local(record)), len(first_officer_with_transfer), "除本地外均为异地。"),
            ],
        ),
    ]

    origin_counts: dict[str, int] = {}
    origin_people: dict[str, list[PersonnelRecord]] = {}
    for record in structure_crew:
        label = map_origin(record.origin)
        origin_counts[label] = origin_counts.get(label, 0) + 1
        origin_people.setdefault(label, []).append(record)
    other_entries = [(label, value) for label, value in origin_counts.items() if label not in ORIGIN_LABELS]
    other_count = sum(value for _, value in other_entries)
    origin_items = [
        make_item(label, origin_counts.get(label, 0), structure_crew_count, "按原单位映射汇总；总队777与777返聘合并到飞行/总队777。")
        for label in ORIGIN_LABELS
    ]
    if other_count:
        detail = "；".join(f"{label}（{format_people(origin_people.get(label, []))}）" for label, _ in other_entries)
        origin_items.append(make_item("其他", other_count, structure_crew_count, f"未映射原单位，需人工核对：{detail}。"))
    sections.append(StatSection("空勤人员原单位情况", f"{structure_crew_count}人", origin_items))

    warnings = []
    for section in sections:
        for item in section.items:
            if item.label == "其他" and item.count:
                warnings.append(f"{section.title}：其他 {item.count} 人，{item.rule}")

    return StatResult(
        total_people=total_people,
        registered_crew_count=registered_crew_count,
        ground_count=ground_count,
        sections=sections,
        warnings=warnings,
    )


def cell_text(cell) -> str:
    return cell.text.strip().replace("\n", " / ")


def set_cell_text(cell, text: str) -> None:
    # Word 单元格里可能有多个 run、隐藏段落或 WPS 产生的残留文本。
    # 直接改第一个 run 会出现“1955”“50%0”这类拼接残留，所以单元格写入必须整体重建正文。
    cell.text = text


def parse_int_text(value: str) -> int | None:
    text = normalize_text(value)
    if not text:
        return None
    match = re.search(r"-?\d+", text)
    return int(match.group(0)) if match else None


def change_text(current: int, previous: int | None) -> str:
    if previous is None:
        return "/"
    delta = current - previous
    if delta > 0:
        return f"新增{delta}人"
    if delta < 0:
        return f"减少{abs(delta)}人"
    return "/"


def find_month_column(table, month: int | None) -> int:
    header = [cell_text(cell) for cell in table.rows[0].cells]
    month_columns = []
    for index, value in enumerate(header):
        match = MONTH_PATTERN.match(value)
        if match:
            month_columns.append((int(match.group(1)), index))
    if not month_columns:
        raise ValueError("表格未识别到月份列。")
    if month is not None:
        for value, index in month_columns:
            if value == month:
                return index
        raise ValueError(f"表格未找到 {month}月 列。")

    empty_candidates = []
    for value, index in month_columns:
        body_values = [cell_text(row.cells[index]) for row in table.rows[1:]]
        if any(item == "" for item in body_values):
            empty_candidates.append((value, index))
    if len(empty_candidates) == 1:
        return empty_candidates[0][1]
    if len(empty_candidates) > 1:
        values = "、".join(f"{value}月" for value, _ in empty_candidates)
        raise ValueError(f"识别到多个存在空值的月份列：{values}，请使用 --month 指定。")
    return month_columns[-1][1]


def column_index_by_header(table, header_text: str) -> int | None:
    header = [cell_text(cell) for cell in table.rows[0].cells]
    for index, value in enumerate(header):
        if value == header_text:
            return index
    return None


def clear_body_columns(table, columns: Iterable[int | None]) -> None:
    clear_columns = sorted({column for column in columns if column is not None})
    for row in table.rows[1:]:
        for column in clear_columns:
            if 0 <= column < len(row.cells):
                set_cell_text(row.cells[column], "")


def section_map(result: StatResult) -> dict[str, StatSection]:
    return {section.title: section for section in result.sections}


def item_map(section: StatSection) -> dict[str, StatItem]:
    return {item.label: item for item in section.items}


def update_group_label(cell, text: str) -> None:
    set_cell_text(cell, text)


def fill_table(
    table,
    section: StatSection,
    month: int | None,
    logs: list[str],
    skip_labels: set[str] | None = None,
    label_column: int = 0,
) -> None:
    skip_labels = skip_labels or set()
    month_col = find_month_column(table, month)
    previous_month_col = month_col - 1 if month_col > 0 else None
    change_col = column_index_by_header(table, "本月变化")
    percent_col = column_index_by_header(table, "本月占比")
    clear_body_columns(table, [month_col, change_col, percent_col])
    items = item_map(section)

    for row in table.rows[1:]:
        label = cell_text(row.cells[label_column]) if label_column < len(row.cells) else ""
        if not label or label in skip_labels:
            continue
        item = items.get(label)
        if item is None:
            logs.append(f"{section.title}：Word 行[{label}]无对应统计项，跳过。")
            continue
        set_cell_text(row.cells[month_col], str(item.count))
        if change_col is not None:
            previous_value = parse_int_text(cell_text(row.cells[previous_month_col])) if previous_month_col is not None else None
            set_cell_text(row.cells[change_col], change_text(item.count, previous_value))
        if percent_col is not None:
            set_cell_text(row.cells[percent_col], item.percent)
    logs.append(f"{section.title}：已写入表格。")


def fill_captain_level_table(table, section: StatSection, month: int | None, logs: list[str]) -> None:
    month_col = find_month_column(table, month)
    previous_month_col = month_col - 1 if month_col > 0 else None
    change_col = column_index_by_header(table, "本月变化")
    percent_col = column_index_by_header(table, "本月占比")
    clear_body_columns(table, [month_col, change_col, percent_col])
    items = item_map(section)
    teacher_count = sum(items[label].count for label in ["C类教员", "B类教员", "A类教员"] if label in items)
    captain_count = section.items[0].denominator - teacher_count

    for row in table.rows[1:]:
        label = cell_text(row.cells[1])
        if not label:
            continue
        if label.endswith("教员"):
            update_group_label(row.cells[0], f"教员 / （{teacher_count}）")
        elif "机长" in label:
            update_group_label(row.cells[0], f"机长 / （{captain_count}）")
        item = items.get(label)
        if item is None:
            logs.append(f"{section.title}：Word 行[{label}]无对应统计项，跳过。")
            continue
        set_cell_text(row.cells[month_col], str(item.count))
        if change_col is not None:
            previous_value = parse_int_text(cell_text(row.cells[previous_month_col])) if previous_month_col is not None else None
            set_cell_text(row.cells[change_col], change_text(item.count, previous_value))
        if percent_col is not None:
            set_cell_text(row.cells[percent_col], item.percent)
    logs.append(f"{section.title}：已写入表格。")


def fill_residence_table(table, section: StatSection, month: int | None, logs: list[str]) -> None:
    month_col = find_month_column(table, month)
    previous_month_col = month_col - 1 if month_col > 0 else None
    change_col = column_index_by_header(table, "本月变化")
    percent_col = column_index_by_header(table, "本月占比")
    clear_body_columns(table, [month_col, change_col, percent_col])
    items = item_map(section)
    captain_denominator = items["机长本地居住"].denominator
    first_officer_denominator = items["副驾驶本地居住"].denominator

    for row in table.rows[1:]:
        group = cell_text(row.cells[0])
        label = cell_text(row.cells[1])
        if "机长" in group:
            update_group_label(row.cells[0], f"机长 / （{captain_denominator}）")
            key = f"机长{label}"
        elif "副驾驶" in group:
            update_group_label(row.cells[0], f"副驾驶（{first_officer_denominator}）")
            key = f"副驾驶{label}"
        else:
            logs.append(f"{section.title}：未识别居住分组[{group}]，跳过。")
            continue
        item = items.get(key)
        if item is None:
            logs.append(f"{section.title}：Word 行[{key}]无对应统计项，跳过。")
            continue
        set_cell_text(row.cells[month_col], str(item.count))
        if change_col is not None:
            previous_value = parse_int_text(cell_text(row.cells[previous_month_col])) if previous_month_col is not None else None
            set_cell_text(row.cells[change_col], change_text(item.count, previous_value))
        if percent_col is not None:
            set_cell_text(row.cells[percent_col], item.percent)
    logs.append(f"{section.title}：已写入表格。")


def fill_origin_table(table, section: StatSection, month: int | None, logs: list[str]) -> None:
    month_col = find_month_column(table, month)
    previous_month_col = month_col - 1 if month_col > 0 else None
    change_col = column_index_by_header(table, "本月变化")
    percent_col = column_index_by_header(table, "本月占比")
    clear_body_columns(table, [month_col, change_col, percent_col])
    items = item_map(section)
    origin_alias = {
        "777": "飞行/总队 777",
        "737": "飞行/总队 737",
        "320": "飞行/总队 320",
        "909": "飞行/总队 909",
    }

    for row in table.rows[1:]:
        left = cell_text(row.cells[0])
        right = cell_text(row.cells[1])
        label = origin_alias.get(right) or right or left
        item = items.get(label)
        if item is None:
            logs.append(f"{section.title}：Word 行[{label}]无对应统计项，跳过。")
            continue
        set_cell_text(row.cells[month_col], str(item.count))
        if change_col is not None:
            previous_value = parse_int_text(cell_text(row.cells[previous_month_col])) if previous_month_col is not None else None
            set_cell_text(row.cells[change_col], change_text(item.count, previous_value))
        if percent_col is not None:
            set_cell_text(row.cells[percent_col], item.percent)
    logs.append(f"{section.title}：已写入表格。")


def fill_docx(docx_path: Path, result: StatResult, output_path: Path, month: int | None, dry_run: bool = False) -> list[str]:
    if Document is None:
        raise RuntimeError("缺少 python-docx，请先运行：pip install openpyxl python-docx")
    document = Document(docx_path)
    if len(document.tables) < 9:
        raise ValueError(f"Word 报告表格数量不足：识别到 {len(document.tables)} 张，至少需要 9 张。")

    logs: list[str] = []
    sections = section_map(result)

    table_sections = [
        "飞行管理人员占比",
        "教员、机长、副驾驶占比",
        "机长含以上各级别占比",
        "机长航线资格占比",
        "机长报务占比",
        "副驾驶级别占比",
        "副驾驶报务占比",
    ]
    for index, title in enumerate(table_sections):
        if title == "机长含以上各级别占比":
            fill_captain_level_table(document.tables[index], sections[title], month, logs)
            continue
        skip_labels = {"其他"} if title == "机长航线资格占比" else set()
        fill_table(document.tables[index], sections[title], month, logs, skip_labels=skip_labels)

    fill_residence_table(document.tables[7], sections["人员居住情况"], month, logs)
    fill_origin_table(document.tables[8], sections["空勤人员原单位情况"], month, logs)

    if dry_run:
        logs.append(f"dry-run：未保存文件，目标路径为 {output_path}")
        return logs

    document.save(output_path)
    logs.append(f"已生成：{output_path}")
    return logs


def default_output_path(docx_path: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return docx_path.with_name(f"{docx_path.stem}_人员结构已填_{timestamp}{docx_path.suffix}")


def choose_file(title: str, filetypes: list[tuple[str, str]]) -> Path | None:
    try:
        import tkinter as tk
        from tkinter import filedialog
    except Exception:
        return None
    root = tk.Tk()
    root.withdraw()
    root.update()
    filename = filedialog.askopenfilename(title=title, filetypes=filetypes)
    root.destroy()
    return Path(filename) if filename else None


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="用人员信息 Excel 填写人员结构报告 docx。")
    parser.add_argument("--excel", type=Path, help="人员信息 Excel 文件路径。")
    parser.add_argument("--docx", type=Path, help="人员结构报告 docx 文件路径。")
    parser.add_argument("--output", type=Path, help="输出 docx 文件路径；默认在原文件旁生成新文件。")
    parser.add_argument("--sheet", help="Excel 工作表名称；默认读取第一个工作表。")
    parser.add_argument("--month", type=int, choices=range(1, 13), metavar="1-12", help="写入月份；不填时自动识别存在空值的月份列，识别不到则写最后一个月份列。")
    parser.add_argument("--dry-run", action="store_true", help="只显示写入计划，不保存 docx。")
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    excel_path = args.excel
    docx_path = args.docx

    if excel_path is None:
        excel_path = choose_file("选择人员信息 Excel", [("Excel 文件", "*.xlsx *.xlsm *.xls"), ("所有文件", "*.*")])
    if docx_path is None:
        docx_path = choose_file("选择人员结构报告 Word", [("Word 文件", "*.docx"), ("所有文件", "*.*")])
    if excel_path is None or docx_path is None:
        print("已取消：未选择 Excel 或 docx。")
        return 1
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel 文件不存在：{excel_path}")
    if not docx_path.exists():
        raise FileNotFoundError(f"docx 文件不存在：{docx_path}")

    output_path = args.output or default_output_path(docx_path)
    records = parse_excel(excel_path, args.sheet)
    result = calculate(records)
    logs = fill_docx(docx_path, result, output_path, args.month, args.dry_run)

    print(f"已读取人员记录：{len(records)}")
    print(f"总人数：{result.total_people}，已注册空勤：{result.registered_crew_count}，地面人员：{result.ground_count}")
    if result.warnings:
        print("需要核对：")
        for warning in result.warnings:
            print(f"  - {warning}")
    print("写入结果：")
    for log in logs:
        print(f"  - {log}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main(sys.argv[1:]))
    except Exception as exc:
        print(f"错误：{exc}", file=sys.stderr)
        raise SystemExit(1)
